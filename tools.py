from playwright.async_api import async_playwright
from langchain_community.agent_toolkits import PlayWrightBrowserToolkit
from dotenv import load_dotenv
import os
import requests
from langchain.agents import Tool
from langchain_core.tools.structured import StructuredTool
from langchain_community.agent_toolkits import FileManagementToolkit
from langchain_community.tools.wikipedia.tool import WikipediaQueryRun
from langchain_experimental.tools import PythonREPLTool
from langchain_community.utilities import GoogleSerperAPIWrapper
from langchain_community.utilities.wikipedia import WikipediaAPIWrapper
from utils import safe_tool
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from typing import Dict, List
from pydantic import EmailStr, ValidationError, BaseModel, Field
import logging
import google.oauth2.service_account as sa
from markdown_pdf import MarkdownPdf, Section
from docx import Document

load_dotenv(override=True)
pushover_token = os.getenv("PUSHOVER_TOKEN")
pushover_user = os.getenv("PUSHOVER_USER")
pushover_url = "https://api.pushover.net/1/messages.json"
serper = GoogleSerperAPIWrapper()

class Mail(BaseModel):
    to_addr: List[EmailStr] = Field(description="Recipient address or list of addresses")
    subject: str = Field(description="Subject")
    body: str = Field(description="Body")

def get_file_tools():
    toolkit = FileManagementToolkit(root_dir="sandbox")
    return toolkit.get_tools()


@safe_tool
def push(text: str):
    """Send a push notification to the user"""
    requests.post(pushover_url, data = {"token": pushover_token, "user": pushover_user, "message": text})
    return {"success": "Push Notification sent successfully"}

@safe_tool
def send_email(to_addr: List[EmailStr], subject: str, body: str) -> Dict[str, str]:
    smtp_server = os.environ.get("SMTP_SERVER")
    smtp_port = 587
    sender_email = os.environ.get("SMTP_EMAIL")
    sender_password = os.environ.get("SMTP_PASSWORD")

    if not smtp_server or not sender_email or not sender_password:
        logging.warning("SMTP configuration missing")
        return {"status": "error", "message": "Missing SMTP configuration"}
    
    if isinstance(to_addr, str):
        recipient_emails = [to_addr]
    else:
        recipient_emails = to_addr

    for r in recipient_emails:
        try:
            EmailStr._validate(r)
        except ValidationError:
            logging.error(f"Invalid recipient email: {r}")
            return {"status": "error", "message": f"Invalid recipient email: {r}"}

    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(sender_email, sender_password)

            msg = MIMEMultipart("alternative")
            msg["Subject"] = subject
            msg["From"] = sender_email
            msg["To"] = ", ".join(recipient_emails)
            msg.attach(MIMEText(body, "html"))

            server.sendmail(sender_email, recipient_emails, msg.as_string())
            logging.info(f"Emails sent successfully to {', '.join(recipient_emails)}")
        return {"status": "success", "message": f"Emails sent to {', '.join(recipient_emails)}"}
    except Exception as e:
        logging.error(f"Failed to send email: {e}")
        return {"status": "error", "message": f"Failed to send email: {e}"}

@safe_tool
def markdown_to_pdf(file_name: str, output: str = "sandbox/output.pdf"):
    with open(f"sandbox/{file_name}", "r") as f:
        markdown_content = f.read()
    pdf = MarkdownPdf()
    pdf.add_section(Section(markdown_content))
    pdf.save(output)
    return {"success": True, "output_path": output}

@safe_tool
def read_docx(file_path: str) -> Dict[str, str]:
    """Read a DOCX file and extract all text content"""
    try:
        doc = Document(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        content = "\n".join(full_text)
        return {"success": True, "content": content, "file_path": file_path}
    except Exception as e:
        return {"success": False, "error": str(e), "file_path": file_path}


async def playwright_tools():
    playwright = await async_playwright().start()
    browser = await playwright.chromium.launch(headless=False)
    toolkit = PlayWrightBrowserToolkit.from_browser(async_browser=browser)
    return toolkit.get_tools(), browser, playwright

async def other_tools():
    push_tool = Tool(name="notify_user", func=push, description="Use this tool when you want to send a push notification to the user")
    email_tool = StructuredTool(name="send_email", func=send_email, description="Use this tool when you want to send an email", args_schema=Mail)
    markdown_to_pdf_tool = Tool(name="markdown_to_pdf", func=markdown_to_pdf, description="Use this tool when you want to convert Markdown(.md) file to PDF")
    docx_reader_tool = Tool(name="read_docx", func=read_docx, description="Use this tool when you want to read and extract text content from a DOCX file")
    file_tools = get_file_tools()

    tool_search =Tool(
        name="search",
        func=serper.run,
        description="Use this tool when you want to get the results of an online web search"
    )

    wikipedia = WikipediaAPIWrapper()
    wiki_tool = WikipediaQueryRun(api_wrapper=wikipedia)

    python_repl = PythonREPLTool()
    
    return file_tools + [push_tool, tool_search, python_repl, wiki_tool, email_tool, markdown_to_pdf_tool, docx_reader_tool]

